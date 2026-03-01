#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Script de génération de documentation non technique pour l'application de Gestion Comptabilité
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_heading_with_color(doc, text, level, color=(0, 51, 102)):
    """Ajoute un titre avec une couleur personnalisée"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(*color)
    return heading

def add_horizontal_line(paragraph):
    """Ajoute une ligne horizontale sous un paragraphe"""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '003366')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_documentation():
    """Crée le document Word de documentation"""
    doc = Document()
    
    # Configuration du document
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # ===== PAGE DE GARDE =====
    title = doc.add_heading('Application de Gestion Comptabilité', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 51, 102)
        run.font.size = Pt(28)
    
    subtitle = doc.add_paragraph('Guide Utilisateur Non Technique')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(18)
    subtitle.runs[0].font.color.rgb = RGBColor(102, 102, 102)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    version_info = doc.add_paragraph('Version 1.0 - 2026')
    version_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    version_info.runs[0].font.size = Pt(12)
    version_info.runs[0].italic = True
    
    doc.add_page_break()
    
    # ===== TABLE DES MATIÈRES =====
    add_heading_with_color(doc, 'Table des Matières', 1)
    
    toc_items = [
        '1. Introduction',
        '2. Premiers Pas',
        '3. Tableau de Bord',
        '4. Gestion des Comptables',
        '5. Gestion des Clients',
        '6. Gestion des Dossiers',
        '7. Gestion des Honoraires',
        '8. Gestion Fiscale',
        '9. Gestion Juridique',
        '10. Réclamations',
        '11. Exports et Rapports',
        '12. FAQ et Support'
    ]
    
    for item in toc_items:
        p = doc.add_paragraph(item, style='List Number')
        p.runs[0].font.size = Pt(11)
    
    doc.add_page_break()
    
    # ===== 1. INTRODUCTION =====
    add_heading_with_color(doc, '1. Introduction', 1)
    
    add_heading_with_color(doc, '1.1 À propos de l\'application', 2)
    doc.add_paragraph(
        'L\'Application de Gestion Comptabilité est une solution complète conçue pour faciliter '
        'la gestion quotidienne d\'un cabinet comptable. Elle permet de gérer efficacement les '
        'dossiers clients, les honoraires, les obligations fiscales et juridiques, ainsi que '
        'la communication avec les clients.'
    )
    
    add_heading_with_color(doc, '1.2 Fonctionnalités principales', 2)
    features = [
        'Gestion centralisée des dossiers clients (personnes physiques et morales)',
        'Suivi des honoraires et des règlements',
        'Gestion des obligations fiscales (TVA, acomptes, CMIR, dépôt de bilan)',
        'Gestion des documents et événements juridiques',
        'Système de réclamations et notifications',
        'Tableaux de bord et statistiques en temps réel',
        'Exports PDF et Excel personnalisables',
        'Notifications automatiques par email'
    ]
    
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    add_heading_with_color(doc, '1.3 Rôles utilisateurs', 2)
    
    # Tableau des rôles
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid Accent 1'
    
    roles_data = [
        ('Rôle', 'Description'),
        ('Administrateur', 'Accès complet à toutes les fonctionnalités, gestion des utilisateurs et configuration système'),
        ('Manager', 'Supervision des comptables, accès aux rapports et statistiques globales'),
        ('Comptable', 'Gestion des dossiers assignés, saisie des données fiscales et honoraires'),
        ('Client', 'Consultation de ses propres dossiers, documents et honoraires')
    ]
    
    for i, (role, desc) in enumerate(roles_data):
        row = table.rows[i]
        row.cells[0].text = role
        row.cells[1].text = desc
        if i == 0:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    
    doc.add_page_break()
    
    # ===== 2. PREMIERS PAS =====
    add_heading_with_color(doc, '2. Premiers Pas', 1)
    
    add_heading_with_color(doc, '2.1 Connexion à l\'application', 2)
    doc.add_paragraph(
        'Pour accéder à l\'application, suivez ces étapes simples :'
    )
    
    steps = [
        'Ouvrez votre navigateur web (Chrome, Firefox, Edge, Safari)',
        'Accédez à l\'URL de l\'application fournie par votre administrateur',
        'Sur la page de connexion, entrez votre nom d\'utilisateur',
        'Saisissez votre mot de passe',
        'Cliquez sur le bouton "Se connecter"'
    ]
    
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph(f'{i}. {step}')
        p.style = 'List Number'
    
    note = doc.add_paragraph()
    note.add_run('Note : ').bold = True
    note.add_run(
        'Si vous avez oublié votre mot de passe, contactez votre administrateur système. '
        'Pour des raisons de sécurité, les mots de passe ne peuvent pas être récupérés automatiquement.'
    )
    note.runs[1].font.italic = True
    
    add_heading_with_color(doc, '2.2 Interface principale', 2)
    doc.add_paragraph(
        'Une fois connecté, vous accédez au tableau de bord principal qui affiche :'
    )
    
    interface_items = [
        'Menu de navigation en haut de la page',
        'Statistiques et indicateurs clés',
        'Alertes et notifications importantes',
        'Accès rapide aux fonctionnalités principales'
    ]
    
    for item in interface_items:
        doc.add_paragraph(item, style='List Bullet')
    
    add_heading_with_color(doc, '2.3 Navigation dans l\'application', 2)
    doc.add_paragraph(
        'Le menu principal vous permet d\'accéder aux différents modules :'
    )
    
    nav_items = [
        'Tableau de Bord : Vue d\'ensemble et statistiques',
        'Comptables : Gestion des comptables du cabinet',
        'Clients : Liste et gestion des clients',
        'Dossiers : Gestion des dossiers comptables',
        'Honoraires : Suivi des honoraires et règlements',
        'Fiscal : Gestion des obligations fiscales',
        'Juridique : Documents et événements juridiques',
        'Réclamations : Système de réclamations et suivi'
    ]
    
    for item in nav_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 3. TABLEAU DE BORD =====
    add_heading_with_color(doc, '3. Tableau de Bord', 1)
    
    doc.add_paragraph(
        'Le tableau de bord est la page d\'accueil de l\'application. Il fournit une vue d\'ensemble '
        'de l\'activité du cabinet et des informations importantes.'
    )
    
    add_heading_with_color(doc, '3.1 Statistiques générales', 2)
    doc.add_paragraph('Le tableau de bord affiche les statistiques suivantes :')
    
    stats = [
        'Nombre total de comptables actifs',
        'Nombre total de dossiers actifs',
        'Répartition entre personnes physiques (PP) et personnes morales (PM)',
        'Statistiques sur les honoraires (en attente, payés)',
        'Nombre de réclamations urgentes',
        'Alertes sur les impayés et retards'
    ]
    
    for stat in stats:
        doc.add_paragraph(stat, style='List Bullet')
    
    add_heading_with_color(doc, '3.2 Graphiques et visualisations', 2)
    doc.add_paragraph(
        'Des graphiques interactifs vous permettent de visualiser rapidement :'
    )
    
    charts = [
        'Répartition des dossiers par forme juridique',
        'Dossiers par comptable (top 5)',
        'Évolution des créations de dossiers',
        'Statistiques TVA (mensuelle, trimestrielle, exonérée)'
    ]
    
    for chart in charts:
        doc.add_paragraph(chart, style='List Bullet')
    
    add_heading_with_color(doc, '3.3 Alertes et notifications', 2)
    doc.add_paragraph(
        'Le tableau de bord affiche les alertes importantes nécessitant votre attention :'
    )
    
    alerts = [
        'Honoraires impayés ou en retard',
        'Règlements en attente',
        'Réclamations non traitées',
        'Événements juridiques à venir',
        'Documents récents'
    ]
    
    for alert in alerts:
        doc.add_paragraph(alert, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 4. GESTION DES COMPTABLES =====
    add_heading_with_color(doc, '4. Gestion des Comptables', 1)
    
    note = doc.add_paragraph()
    note.add_run('Accès : ').bold = True
    note.add_run('Administrateur et Manager uniquement')
    note.runs[1].font.italic = True
    note.runs[1].font.color.rgb = RGBColor(204, 0, 0)
    
    add_heading_with_color(doc, '4.1 Liste des comptables', 2)
    doc.add_paragraph(
        'La page "Comptables" affiche la liste de tous les comptables actifs du cabinet avec :'
    )
    
    comptable_info = [
        'Nom et prénom',
        'Matricule',
        'Email de contact',
        'Nombre de dossiers assignés (PM et PP)',
        'Statut (actif/inactif)'
    ]
    
    for info in comptable_info:
        doc.add_paragraph(info, style='List Bullet')
    
    add_heading_with_color(doc, '4.2 Ajouter un comptable', 2)
    doc.add_paragraph('Pour ajouter un nouveau comptable :')
    
    add_steps = [
        'Cliquez sur le bouton "Ajouter un comptable"',
        'Remplissez le formulaire avec les informations requises :',
        '  - Nom et prénom',
        '  - Matricule unique',
        '  - Email professionnel',
        '  - Nom d\'utilisateur pour la connexion',
        '  - Mot de passe initial',
        'Cliquez sur "Enregistrer"',
        'Un email de bienvenue sera automatiquement envoyé au comptable avec ses identifiants'
    ]
    
    for step in add_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    add_heading_with_color(doc, '4.3 Modifier un comptable', 2)
    doc.add_paragraph(
        'Pour modifier les informations d\'un comptable, cliquez sur son nom dans la liste, '
        'puis sur le bouton "Modifier". Vous pouvez mettre à jour toutes les informations '
        'sauf le matricule.'
    )
    
    add_heading_with_color(doc, '4.4 Consulter les détails d\'un comptable', 2)
    doc.add_paragraph(
        'En cliquant sur un comptable, vous accédez à sa fiche détaillée qui affiche :'
    )
    
    details = [
        'Informations personnelles et professionnelles',
        'Liste complète des dossiers assignés',
        'Statistiques détaillées (nombre de PM, PP, dossiers forfaitaires, TVA)',
        'Historique d\'activité'
    ]
    
    for detail in details:
        doc.add_paragraph(detail, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 5. GESTION DES CLIENTS =====
    add_heading_with_color(doc, '5. Gestion des Clients', 1)
    
    note = doc.add_paragraph()
    note.add_run('Accès : ').bold = True
    note.add_run('Administrateur uniquement')
    note.runs[1].font.italic = True
    note.runs[1].font.color.rgb = RGBColor(204, 0, 0)
    
    add_heading_with_color(doc, '5.1 Liste des clients', 2)
    doc.add_paragraph(
        'La page "Clients" affiche tous les clients enregistrés dans le système. '
        'Vous pouvez rechercher un client par nom, email ou téléphone.'
    )
    
    add_heading_with_color(doc, '5.2 Ajouter un client', 2)
    doc.add_paragraph('Pour créer un nouveau compte client :')
    
    client_steps = [
        'Cliquez sur "Ajouter un client"',
        'Remplissez les informations requises :',
        '  - Nom et prénom',
        '  - Nom de l\'entreprise',
        '  - Email',
        '  - Téléphone',
        '  - Adresse',
        '  - Nom d\'utilisateur',
        '  - Mot de passe initial',
        'Validez le formulaire',
        'Le client recevra un email avec ses identifiants de connexion'
    ]
    
    for step in client_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    add_heading_with_color(doc, '5.3 Modifier ou supprimer un client', 2)
    doc.add_paragraph(
        'Depuis la liste des clients, vous pouvez modifier les informations d\'un client '
        'ou le supprimer si nécessaire. La suppression d\'un client supprime également '
        'son compte utilisateur associé.'
    )
    
    warning = doc.add_paragraph()
    warning.add_run('Attention : ').bold = True
    warning.runs[0].font.color.rgb = RGBColor(204, 0, 0)
    warning.add_run(
        'La suppression d\'un client est irréversible. Assurez-vous que le client '
        'n\'a pas de dossiers actifs avant de le supprimer.'
    )
    
    doc.add_page_break()
    
    # ===== 6. GESTION DES DOSSIERS =====
    add_heading_with_color(doc, '6. Gestion des Dossiers', 1)
    
    doc.add_paragraph(
        'Les dossiers représentent les clients du cabinet et contiennent toutes les informations '
        'comptables, fiscales et juridiques associées.'
    )
    
    add_heading_with_color(doc, '6.1 Liste des dossiers', 2)
    doc.add_paragraph(
        'La page "Dossiers" affiche tous les dossiers actifs. Vous pouvez filtrer par :'
    )
    
    filters = [
        'Type de structure (Personne Physique / Personne Morale)',
        'Forme juridique (SARL, SA, SAS, EI, AUTO, etc.)',
        'Code dossier (R, F, FS, etc.)',
        'Déclaration TVA (Mensuelle, Trimestrielle, Exonérée)',
        'Comptable traitant (pour administrateurs et managers)'
    ]
    
    for f in filters:
        doc.add_paragraph(f, style='List Bullet')
    
    add_heading_with_color(doc, '6.2 Créer un dossier', 2)
    doc.add_paragraph('Pour créer un nouveau dossier :')
    
    dossier_steps = [
        'Cliquez sur "Nouveau dossier"',
        'Remplissez les informations générales :',
        '  - Dénomination (nom de l\'entreprise ou du client)',
        '  - Abréviation',
        '  - Forme juridique',
        '  - Code dossier (R = Réel, F = Forfaitaire, FS = Forfaitaire Simplifié)',
        '  - ICE (Identifiant Commun de l\'Entreprise)',
        '  - RC (Registre de Commerce)',
        '  - Secteur d\'activité et branche',
        'Sélectionnez le comptable traitant',
        'Sélectionnez le client associé',
        'Configurez les paramètres fiscaux :',
        '  - Type de déclaration TVA',
        '  - Dates de clôture',
        'Enregistrez le dossier'
    ]
    
    for step in dossier_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    add_heading_with_color(doc, '6.3 Consulter un dossier', 2)
    doc.add_paragraph(
        'En cliquant sur un dossier, vous accédez à sa fiche complète qui regroupe :'
    )
    
    dossier_sections = [
        'Informations générales du dossier',
        'Honoraires associés',
        'Suivis fiscaux (TVA, acomptes, CMIR, dépôt de bilan)',
        'Documents et événements juridiques',
        'Réclamations liées au dossier'
    ]
    
    for section in dossier_sections:
        doc.add_paragraph(section, style='List Bullet')
    
    add_heading_with_color(doc, '6.4 Modifier ou archiver un dossier', 2)
    doc.add_paragraph(
        'Vous pouvez modifier les informations d\'un dossier à tout moment. '
        'Pour archiver un dossier (le rendre inactif), décochez la case "Actif" '
        'dans le formulaire de modification.'
    )
    
    doc.add_page_break()
    
    # ===== 7. GESTION DES HONORAIRES =====
    add_heading_with_color(doc, '7. Gestion des Honoraires', 1)
    
    doc.add_paragraph(
        'Le module Honoraires permet de gérer les facturations et les règlements des clients.'
    )
    
    add_heading_with_color(doc, '7.1 Types d\'honoraires', 2)
    doc.add_paragraph('L\'application gère deux types d\'honoraires :')
    
    types_honoraires = [
        'Honoraires standards : Facturations régulières pour les prestations comptables',
        'Honoraires PV : Honoraires liés aux procès-verbaux et formalités juridiques'
    ]
    
    for t in types_honoraires:
        doc.add_paragraph(t, style='List Bullet')
    
    add_heading_with_color(doc, '7.2 Créer un honoraire', 2)
    doc.add_paragraph('Pour créer un nouvel honoraire :')
    
    honoraire_steps = [
        'Accédez au dossier concerné',
        'Dans la section "Honoraires", cliquez sur "Ajouter un honoraire"',
        'Remplissez les informations :',
        '  - Montant HT',
        '  - Taux de TVA',
        '  - Date d\'émission',
        '  - Date d\'échéance',
        '  - Description de la prestation',
        'Le montant TTC est calculé automatiquement',
        'Enregistrez l\'honoraire'
    ]
    
    for step in honoraire_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    add_heading_with_color(doc, '7.3 Statuts de règlement', 2)
    doc.add_paragraph('Les honoraires peuvent avoir les statuts suivants :')
    
    # Tableau des statuts
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    statuts_data = [
        ('Statut', 'Description'),
        ('EN_ATTENTE', 'Honoraire émis, en attente de paiement'),
        ('PAYE', 'Honoraire intégralement réglé'),
        ('EN_RETARD', 'Date d\'échéance dépassée, paiement non reçu')
    ]
    
    for i, (statut, desc) in enumerate(statuts_data):
        row = table.rows[i]
        row.cells[0].text = statut
        row.cells[1].text = desc
        if i == 0:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    
    doc.add_paragraph()
    
    add_heading_with_color(doc, '7.4 Enregistrer un règlement', 2)
    doc.add_paragraph(
        'Lorsqu\'un client effectue un paiement, vous pouvez enregistrer le règlement :'
    )
    
    reglement_steps = [
        'Accédez à l\'honoraire concerné',
        'Cliquez sur "Enregistrer un règlement"',
        'Saisissez le montant payé',
        'Indiquez la date de paiement',
        'Sélectionnez le mode de paiement (chèque, virement, espèces, etc.)',
        'Ajoutez une référence si nécessaire',
        'Validez le règlement'
    ]
    
    for step in reglement_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    doc.add_paragraph(
        'Le statut de l\'honoraire sera automatiquement mis à jour en fonction des règlements enregistrés.'
    )
    
    doc.add_page_break()
    
    # ===== 8. GESTION FISCALE =====
    add_heading_with_color(doc, '8. Gestion Fiscale', 1)
    
    doc.add_paragraph(
        'Le module Fiscal permet de gérer toutes les obligations fiscales des dossiers.'
    )
    
    add_heading_with_color(doc, '8.1 Suivi de la TVA', 2)
    doc.add_paragraph(
        'Pour les dossiers soumis à la TVA (mensuelle ou trimestrielle), vous pouvez créer '
        'des suivis TVA pour chaque période :'
    )
    
    tva_steps = [
        'Accédez au dossier concerné',
        'Dans la section "Fiscal", cliquez sur "Nouveau suivi TVA"',
        'Sélectionnez la période (mois ou trimestre)',
        'Saisissez les montants :',
        '  - TVA collectée',
        '  - TVA déductible',
        '  - TVA à payer (calculée automatiquement)',
        'Indiquez la date de dépôt',
        'Ajoutez des observations si nécessaire',
        'Enregistrez le suivi'
    ]
    
    for step in tva_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    add_heading_with_color(doc, '8.2 Acomptes provisionnels', 2)
    doc.add_paragraph(
        'Pour les dossiers soumis aux acomptes provisionnels (IS, IR), vous pouvez créer '
        'et suivre les acomptes trimestriels.'
    )
    
    add_heading_with_color(doc, '8.3 CMIR (Contribution Minimale sur le Revenu)', 2)
    doc.add_paragraph(
        'Gérez les déclarations et paiements de la CMIR pour les dossiers concernés.'
    )
    
    add_heading_with_color(doc, '8.4 Dépôt de bilan', 2)
    doc.add_paragraph(
        'Enregistrez les dépôts de bilan annuels avec les dates de clôture et de dépôt.'
    )
    
    add_heading_with_color(doc, '8.5 Suivi forfaitaire', 2)
    doc.add_paragraph(
        'Pour les dossiers au régime forfaitaire, gérez les déclarations annuelles simplifiées.'
    )
    
    doc.add_page_break()
    
    # ===== 9. GESTION JURIDIQUE =====
    add_heading_with_color(doc, '9. Gestion Juridique', 1)
    
    doc.add_paragraph(
        'Le module Juridique permet de gérer les documents et événements juridiques des dossiers.'
    )
    
    add_heading_with_color(doc, '9.1 Documents juridiques', 2)
    doc.add_paragraph(
        'Vous pouvez télécharger et archiver tous les documents juridiques importants :'
    )
    
    doc_types = [
        'Statuts de société',
        'Procès-verbaux d\'assemblée',
        'Modifications statutaires',
        'Contrats',
        'Certificats d\'immatriculation',
        'Autres documents officiels'
    ]
    
    for doc_type in doc_types:
        doc.add_paragraph(doc_type, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph('Pour ajouter un document :')
    
    doc_steps = [
        'Accédez au dossier concerné',
        'Dans la section "Juridique", cliquez sur "Ajouter un document"',
        'Sélectionnez le type de document',
        'Téléchargez le fichier (PDF, Word, etc.)',
        'Indiquez la date du document',
        'Ajoutez une description',
        'Enregistrez'
    ]
    
    for step in doc_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    add_heading_with_color(doc, '9.2 Événements juridiques', 2)
    doc.add_paragraph(
        'Créez des rappels pour les événements juridiques importants :'
    )
    
    events = [
        'Assemblées générales ordinaires (AGO)',
        'Assemblées générales extraordinaires (AGE)',
        'Dates de renouvellement de mandats',
        'Échéances de formalités',
        'Autres événements à ne pas manquer'
    ]
    
    for event in events:
        doc.add_paragraph(event, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Les événements à venir sont affichés sur le tableau de bord pour ne rien oublier.'
    )
    
    doc.add_page_break()
    
    # ===== 10. RÉCLAMATIONS =====
    add_heading_with_color(doc, '10. Réclamations', 1)
    
    doc.add_paragraph(
        'Le système de réclamations permet une communication efficace entre les clients, '
        'les comptables et l\'administration.'
    )
    
    add_heading_with_color(doc, '10.1 Créer une réclamation', 2)
    doc.add_paragraph('Pour créer une nouvelle réclamation :')
    
    reclamation_steps = [
        'Accédez au dossier concerné (ou à la page Réclamations)',
        'Cliquez sur "Nouvelle réclamation"',
        'Remplissez le formulaire :',
        '  - Objet de la réclamation',
        '  - Description détaillée',
        '  - Niveau de priorité (Normal, Urgent)',
        '  - Destinataire (comptable ou administrateur)',
        'Envoyez la réclamation'
    ]
    
    for step in reclamation_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    doc.add_paragraph(
        'Le destinataire recevra une notification par email et pourra consulter la réclamation '
        'dans son tableau de bord.'
    )
    
    add_heading_with_color(doc, '10.2 Suivre une réclamation', 2)
    doc.add_paragraph(
        'Vous pouvez consulter l\'état de vos réclamations dans la section "Réclamations". '
        'Les statuts possibles sont :'
    )
    
    # Tableau des statuts de réclamation
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    reclamation_statuts = [
        ('Statut', 'Description'),
        ('EN_ATTENTE', 'Réclamation créée, en attente de traitement'),
        ('EN_COURS', 'Réclamation en cours de traitement'),
        ('RESOLUE', 'Réclamation traitée et résolue')
    ]
    
    for i, (statut, desc) in enumerate(reclamation_statuts):
        row = table.rows[i]
        row.cells[0].text = statut
        row.cells[1].text = desc
        if i == 0:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    
    doc.add_page_break()
    
    # ===== 11. EXPORTS ET RAPPORTS =====
    add_heading_with_color(doc, '11. Exports et Rapports', 1)
    
    doc.add_paragraph(
        'L\'application permet d\'exporter les données sous différents formats pour faciliter '
        'l\'analyse et le partage d\'informations.'
    )
    
    add_heading_with_color(doc, '11.1 Exports Excel', 2)
    doc.add_paragraph('Vous pouvez exporter en Excel :')
    
    excel_exports = [
        'Liste des dossiers avec filtres',
        'Liste des comptables',
        'Honoraires par période',
        'Suivis TVA',
        'Statistiques globales'
    ]
    
    for export in excel_exports:
        doc.add_paragraph(export, style='List Bullet')
    
    doc.add_paragraph()
    doc.add_paragraph(
        'Pour exporter, cliquez sur le bouton "Exporter Excel" disponible sur les pages de listes.'
    )
    
    add_heading_with_color(doc, '11.2 Exports PDF', 2)
    doc.add_paragraph('Les exports PDF sont disponibles pour :')
    
    pdf_exports = [
        'Fiches dossiers complètes',
        'Honoraires et factures',
        'Rapports de synthèse',
        'Documents juridiques'
    ]
    
    for export in pdf_exports:
        doc.add_paragraph(export, style='List Bullet')
    
    add_heading_with_color(doc, '11.3 Rapports statistiques', 2)
    doc.add_paragraph(
        'Les administrateurs ont accès à des rapports statistiques avancés :'
    )
    
    reports = [
        'Évolution des dossiers par mois et par année',
        'Répartition par secteur d\'activité',
        'Répartition par branche',
        'Statistiques par comptable',
        'Analyse des honoraires'
    ]
    
    for report in reports:
        doc.add_paragraph(report, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== 12. FAQ ET SUPPORT =====
    add_heading_with_color(doc, '12. FAQ et Support', 1)
    
    add_heading_with_color(doc, '12.1 Questions fréquentes', 2)
    
    # FAQ 1
    q1 = doc.add_paragraph()
    q1.add_run('Q : J\'ai oublié mon mot de passe, que faire ?').bold = True
    doc.add_paragraph(
        'R : Contactez votre administrateur système qui pourra réinitialiser votre mot de passe. '
        'Pour des raisons de sécurité, les mots de passe ne peuvent pas être récupérés automatiquement.'
    )
    
    # FAQ 2
    q2 = doc.add_paragraph()
    q2.add_run('Q : Comment puis-je voir uniquement mes dossiers ?').bold = True
    doc.add_paragraph(
        'R : Si vous êtes comptable, l\'application affiche automatiquement uniquement les dossiers '
        'qui vous sont assignés. Si vous êtes client, vous ne voyez que vos propres dossiers.'
    )
    
    # FAQ 3
    q3 = doc.add_paragraph()
    q3.add_run('Q : Puis-je modifier un honoraire déjà payé ?').bold = True
    doc.add_paragraph(
        'R : Non, pour des raisons de traçabilité comptable, les honoraires payés ne peuvent pas '
        'être modifiés. Si une correction est nécessaire, contactez votre administrateur.'
    )
    
    # FAQ 4
    q4 = doc.add_paragraph()
    q4.add_run('Q : Comment recevoir des notifications par email ?').bold = True
    doc.add_paragraph(
        'R : Les notifications sont envoyées automatiquement pour les événements importants '
        '(nouvelles réclamations, échéances, etc.). Assurez-vous que votre adresse email '
        'est correcte dans votre profil.'
    )
    
    # FAQ 5
    q5 = doc.add_paragraph()
    q5.add_run('Q : Puis-je exporter toutes les données d\'un dossier ?').bold = True
    doc.add_paragraph(
        'R : Oui, depuis la fiche d\'un dossier, vous pouvez exporter un PDF complet contenant '
        'toutes les informations (honoraires, fiscal, juridique).'
    )
    
    add_heading_with_color(doc, '12.2 Support technique', 2)
    doc.add_paragraph(
        'Pour toute assistance technique ou question non couverte par ce guide, '
        'contactez votre administrateur système ou l\'équipe de support.'
    )
    
    support_info = doc.add_paragraph()
    support_info.add_run('Email : ').bold = True
    support_info.add_run('admin@cabinet-comptable.com')
    
    doc.add_paragraph()
    
    add_heading_with_color(doc, '12.3 Bonnes pratiques', 2)
    doc.add_paragraph('Pour une utilisation optimale de l\'application :')
    
    best_practices = [
        'Mettez à jour régulièrement les informations des dossiers',
        'Enregistrez les règlements dès réception des paiements',
        'Créez des réclamations pour toute question ou problème',
        'Consultez le tableau de bord quotidiennement pour les alertes',
        'Exportez régulièrement vos données pour archivage',
        'Utilisez des mots de passe sécurisés et ne les partagez jamais',
        'Déconnectez-vous après chaque session, surtout sur ordinateur partagé'
    ]
    
    for practice in best_practices:
        doc.add_paragraph(practice, style='List Bullet')
    
    doc.add_page_break()
    
    # ===== CONCLUSION =====
    add_heading_with_color(doc, 'Conclusion', 1)
    
    doc.add_paragraph(
        'Ce guide utilisateur vous a présenté les principales fonctionnalités de l\'Application '
        'de Gestion Comptabilité. L\'interface intuitive et les fonctionnalités complètes vous '
        'permettront de gérer efficacement votre activité comptable au quotidien.'
    )
    
    doc.add_paragraph()
    
    doc.add_paragraph(
        'N\'hésitez pas à explorer l\'application et à contacter le support en cas de besoin. '
        'Des mises à jour régulières apporteront de nouvelles fonctionnalités pour améliorer '
        'continuellement votre expérience.'
    )
    
    doc.add_paragraph()
    
    footer = doc.add_paragraph('Merci d\'utiliser l\'Application de Gestion Comptabilité !')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.bold = True
    footer.runs[0].font.size = Pt(14)
    footer.runs[0].font.color.rgb = RGBColor(0, 51, 102)
    
    # Sauvegarder le document
    output_path = os.path.join(
        os.path.dirname(__file__),
        'Documentation_Utilisateur_Gestion_Comptabilite.docx'
    )
    doc.save(output_path)
    print(f"✅ Documentation créée avec succès : {output_path}")
    return output_path

if __name__ == '__main__':
    create_documentation()
